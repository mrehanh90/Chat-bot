from pathlib import Path
from datetime import date

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / 'WhatsApp_AI_Assistant_Project_Documentation_Modular.docx'
ASSETS = ROOT / '.doc-assets'
ASSETS.mkdir(exist_ok=True)

NAVY = '16324F'
BLUE = '2E74B5'
LIGHT_BLUE = 'E8EEF5'
LIGHT_GRAY = 'F2F4F7'
TEXT = '202124'
MUTED = '5B6573'
GREEN = '2D6A4F'
AMBER = '7A5A00'


def font(size=12, bold=False):
    try:
        return ImageFont.truetype('C:/Windows/Fonts/arial.ttf', size)
    except OSError:
        return ImageFont.load_default()


def diagram(path, title, nodes, edges, width=1600, height=760):
    image = Image.new('RGB', (width, height), 'white')
    d = ImageDraw.Draw(image)
    title_font = font(36, True)
    label_font = font(25, True)
    small_font = font(21)
    d.text((55, 30), title, fill='#' + NAVY, font=title_font)
    for edge in edges:
        if len(edge) == 4:
            x1, y1, x2, y2 = edge
        else:
            (x1, y1), (x2, y2) = edge
        d.line((x1, y1, x2, y2), fill='#' + BLUE, width=6)
        # arrowhead
        if abs(x2 - x1) >= abs(y2 - y1):
            direction = 1 if x2 > x1 else -1
            d.polygon([(x2, y2), (x2 - 20 * direction, y2 - 12), (x2 - 20 * direction, y2 + 12)], fill='#' + BLUE)
        else:
            direction = 1 if y2 > y1 else -1
            d.polygon([(x2, y2), (x2 - 12, y2 - 20 * direction), (x2 + 12, y2 - 20 * direction)], fill='#' + BLUE)
    for node in nodes:
        x, y, w, h, heading, lines, fill = node
        d.rounded_rectangle((x, y, x + w, y + h), radius=22, fill='#' + fill, outline='#' + BLUE, width=3)
        d.text((x + 20, y + 18), heading, fill='#' + NAVY, font=label_font)
        ty = y + 58
        for line in lines:
            d.text((x + 20, ty), line, fill='#' + TEXT, font=small_font)
            ty += 29
    image.save(path)


def create_diagrams():
    diagram(ASSETS / 'architecture.png', 'Multi-session architecture', [
        (65, 145, 280, 150, 'WhatsApp number A', ['Owner-linked account'], LIGHT_BLUE),
        (65, 420, 280, 150, 'WhatsApp number B', ['Owner-linked account'], LIGHT_BLUE),
        (500, 145, 330, 150, 'Session Manager', ['Independent WASocket', 'per userId'], 'DDEBF7'),
        (500, 420, 330, 150, 'Session Manager', ['Independent WASocket', 'per userId'], 'DDEBF7'),
        (1050, 75, 390, 150, 'Isolated storage', ['sessions/userId/', 'SQLite + credentials'], LIGHT_GRAY),
        (1050, 370, 390, 150, 'Connected services', ['OpenRouter', 'Google Calendar'], LIGHT_GRAY),
    ], [
        ((345, 220, 500, 220)), ((345, 495, 500, 495)),
        ((830, 220, 1050, 150)), ((830, 265, 1050, 430)),
        ((830, 495, 1050, 480)), ((830, 450, 1050, 445)),
    ])
    diagram(ASSETS / 'message_flow.png', 'Incoming message workflow', [
        (55, 180, 245, 135, 'Incoming message', ['Text, caption,', 'or voice note'], LIGHT_BLUE),
        (370, 180, 250, 135, 'Normalize', ['Extract text or', 'transcribe audio'], 'DDEBF7'),
        (700, 80, 300, 135, 'Store + detect', ['Contact log; meeting,', 'time, location'], LIGHT_GRAY),
        (700, 340, 300, 135, 'Reply policy', ['Flood guard; Away', 'or Advisor profile'], LIGHT_GRAY),
        (1090, 80, 385, 135, 'Owner alert / task', ['Structured owner alert;', 'save task when extracted'], 'F9E6B3'),
        (1090, 340, 385, 135, 'AI response', ['General AI or live', 'web-grounded lookup'], 'DDEBF7'),
        (1090, 565, 385, 110, 'WhatsApp reply', ['Reply to the sender'], 'DDEBF7'),
    ], [
        ((300, 247, 370, 247)), ((620, 230, 700, 147)), ((620, 270, 700, 407)),
        ((1000, 147, 1090, 147)), ((1000, 407, 1090, 407)), ((1280, 475, 1280, 565)),
    ], height=760)
    diagram(ASSETS / 'calendar_flow.png', 'Meeting and Google Calendar workflow', [
        (45, 205, 270, 135, 'Meeting message', ['Date/time in', 'incoming text'], LIGHT_BLUE),
        (400, 205, 300, 135, 'Task extraction', ['Create local task', 'with scheduled time'], 'DDEBF7'),
        (790, 95, 310, 135, 'Calendar connected?', ['Per-user Google', 'OAuth token'], LIGHT_GRAY),
        (790, 390, 310, 135, 'No valid date/time', ['Keep task locally;', 'manual review'], 'F9E6B3'),
        (1190, 95, 340, 135, 'Create event', ['Primary Calendar;', 'popup reminder'], 'DDEBF7'),
        (1190, 390, 340, 135, 'Owner notification', ['Calendar event link', 'in self-chat'], 'DDEBF7'),
    ], [
        ((315, 272, 400, 272)), ((700, 245, 790, 162)), ((700, 300, 790, 457)),
        ((1100, 162, 1190, 162)), ((1360, 230, 1360, 390)),
    ], height=680)
    diagram(ASSETS / 'pairing_flow.png', 'Linking a new WhatsApp assistant', [
        (55, 220, 270, 125, 'Platform / terminal', ['Register a userId', 'and phone number'], LIGHT_BLUE),
        (420, 220, 300, 125, 'Pairing code', ['Baileys requests a', 'short-lived code'], 'DDEBF7'),
        (815, 100, 300, 125, 'Phone action', ['WhatsApp Linked Devices', 'Link with phone number'], LIGHT_GRAY),
        (815, 410, 300, 125, 'Alternative', ['Scan a QR code', 'on the same device'], LIGHT_GRAY),
        (1210, 220, 320, 125, 'Connected session', ['Runs with npm start;', 'auto replies enabled'], 'DDEBF7'),
    ], [
        ((325, 282, 420, 282)), ((720, 255, 815, 162)), ((720, 310, 815, 472)),
        ((1115, 162, 1210, 255)), ((1115, 472, 1210, 310)),
    ], height=650)


def set_cell_shading(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tc_pr.append(shd)


def set_cell_width(cell, width_twips):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn('w:tcW'))
    if tc_w is None:
        tc_w = OxmlElement('w:tcW')
        tc_pr.append(tc_w)
    tc_w.set(qn('w:w'), str(width_twips))
    tc_w.set(qn('w:type'), 'dxa')


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in('w:tblW')
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), '9360')
    tbl_w.set(qn('w:type'), 'dxa')
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)


def set_run(run, size=11, color=TEXT, bold=False, italic=False):
    run.font.name = 'Calibri'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_paragraph(doc, text='', style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run(r, 11, TEXT, True)
        r = p.add_run(text[len(bold_prefix):])
        set_run(r)
    else:
        r = p.add_run(text)
        set_run(r)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_run(r)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_run(r)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_geometry(table, widths)
    hdr = table.rows[0].cells
    for i, label in enumerate(headers):
        set_cell_shading(hdr[i], LIGHT_BLUE)
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(label)
        set_run(r, 10, NAVY, True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            r = p.add_run(value)
            set_run(r, 10, TEXT)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return table


def add_callout(doc, label, text, color=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, color)
    p = cell.paragraphs[0]
    r = p.add_run(label + ' ')
    set_run(r, 10, NAVY, True)
    r = p.add_run(text)
    set_run(r, 10, TEXT)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_figure(doc, image_path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(6.35))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(9)
    r = p.add_run(caption)
    set_run(r, 9, MUTED, italic=True)


def chapter(doc, title):
    """Start every major module on a fresh page for printed-document updates."""
    doc.add_page_break()
    doc.add_heading(title, level=1)


def setup_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [
        ('Heading 1', 16, BLUE, 18, 10),
        ('Heading 2', 13, BLUE, 14, 7),
        ('Heading 3', 12, NAVY, 10, 5),
    ]:
        style = styles[name]
        style.font.name = 'Calibri'
        style._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
        style._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run('WhatsApp AI Assistant | Project Documentation')
    set_run(r, 8.5, MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run('Internal reference guide | Keep credentials and session data private')
    set_run(r, 8.5, MUTED)


def build_document():
    create_diagrams()
    doc = Document()
    setup_document(doc)

    # Cover
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(82)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('WHATSAPP AI ASSISTANT')
    set_run(r, 29, NAVY, True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Complete Project Documentation and Operating Guide')
    set_run(r, 16, BLUE)
    p.paragraph_format.space_after = Pt(20)
    add_callout(doc, 'Purpose:', 'A practical guide for operating, configuring, securing, and extending the multi-session WhatsApp assistant.', LIGHT_BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(42)
    r = p.add_run('Version: Current local implementation\nPrepared: ' + date.today().isoformat())
    set_run(r, 11, MUTED)
    doc.add_page_break()
    doc.add_heading('Document control and amendment log', level=1)
    add_paragraph(doc, 'Use this page to record every printed-document change. When a feature changes, update the affected module, print only that module again, and insert it behind the matching numbered tab.')
    add_table(doc, ['Field', 'Current value'], [
        ['Document title', 'WhatsApp AI Assistant - Project Documentation and Operating Guide'],
        ['Document owner', 'Project administrator'],
        ['Current edition', 'Modular edition'],
        ['Revision date', date.today().isoformat()],
        ['Print update rule', 'Replace or insert only the changed numbered module.'],
    ], [2500, 6860])
    doc.add_heading('Amendment record', level=2)
    add_table(doc, ['Date', 'Module / page inserted', 'Change summary', 'Updated by'], [
        ['', '', '', ''], ['', '', '', ''], ['', '', '', ''], ['', '', '', ''],
    ], [1600, 2550, 3800, 1410])
    add_callout(doc, 'Print-management rule:', 'Keep the cover page and this amendment log at the front. Every major topic starts on a fresh page so a revised module can be reprinted and inserted without reprinting the full manual.', LIGHT_BLUE)

    chapter(doc, 'Contents')
    add_numbered(doc, [
        'Project overview and scope', 'Architecture and data isolation', 'Installation and configuration',
        'WhatsApp registration and phone pairing', 'Modes and message behavior',
        'Meeting, tasks, location, and Calendar automation', 'AI, live information, and voice notes',
        'Operations, monitoring, and troubleshooting', 'Security, privacy, and limitations',
        'Future web-platform roadmap',
    ])

    chapter(doc, '1. Project overview and scope')
    add_paragraph(doc, 'WhatsApp AI Assistant is a Node.js service built with Baileys and OpenRouter. It connects one or more WhatsApp accounts, receives incoming messages, and applies an assistant profile for each independent session.')
    add_callout(doc, 'Important:', 'Baileys is an unofficial WhatsApp Web client. Use only WhatsApp accounts you own or are authorized to manage. Automated use can be restricted by WhatsApp and must comply with applicable laws and platform terms.', 'F9E6B3')
    doc.add_heading('What the system does', level=2)
    add_bullets(doc, [
        'Runs separate WhatsApp sessions under sessions/<userId>/.',
        'Replies to incoming messages when Away Mode is enabled.',
        'Supports Advisor Mode for natural English and Roman Urdu conversations and supportive everyday advice.',
        'Answers current-information questions through OpenRouter web search when available.',
        'Transcribes push-to-talk voice notes through the configured transcription provider.',
        'Records eligible contact metadata and explicitly shared locations.',
        'Forwards meeting, time, and location alerts to the account owner self-chat.',
        'Stores extracted tasks and creates Google Calendar events for valid detected meetings when Calendar is connected.',
    ])

    chapter(doc, '2. Architecture and data isolation')
    add_paragraph(doc, 'Each userId has an independent Baileys authentication directory, local SQLite database, owner identity, reply-rate state, task list, and Google Calendar token. A connection or logout failure in one session does not stop the other sessions.')
    add_figure(doc, ASSETS / 'architecture.png', 'Figure 1. Each WhatsApp account has isolated session credentials, data, and integrations.')
    add_table(doc, ['Component', 'Responsibility', 'Storage / boundary'], [
        ['index.js', 'Session manager, WhatsApp events, commands, reply flow, alerts', 'One WASocket per userId'],
        ['src/openrouterClient.js', 'Structured AI replies, live web queries, voice transcription', 'Uses OpenRouter API key'],
        ['src/dataStore.js', 'SQLite contacts, tasks, reply times, message dedupe, Calendar mappings', 'sessions/<userId>/assistant.sqlite'],
        ['src/googleCalendar.js', 'Google OAuth and Calendar event creation', 'Per-user encrypted token'],
        ['src/sessionStore.js', 'Session paths and session metadata', 'sessions/<userId>/session.json'],
    ], [1800, 4100, 3460])

    chapter(doc, '3. Installation and configuration')
    doc.add_heading('Requirements', level=2)
    add_bullets(doc, ['Node.js 22.5 or later.', 'A WhatsApp account for every session.', 'An OpenRouter API key.', 'OpenRouter balance for dependable web search and voice transcription.', 'Google Cloud OAuth credentials only if Google Calendar automation is required.'])
    doc.add_heading('Install', level=2)
    add_paragraph(doc, 'From the project directory:', bold_prefix='From the project directory:')
    add_callout(doc, 'Command:', 'npm.cmd install\nCopy-Item .env.example .env', LIGHT_GRAY)
    doc.add_heading('Environment variables', level=2)
    add_table(doc, ['Variable', 'Purpose', 'Recommended / notes'], [
        ['OPENROUTER_API_KEY', 'Authorizes AI, live search, and audio requests', 'Required. Never commit or share it.'],
        ['OPENROUTER_MODEL', 'General replies and task extraction', 'Free models may be rate limited.'],
        ['OPENROUTER_LIVE_MODEL', 'Current/live questions with web search', 'Requires model availability and usually credit.'],
        ['OPENROUTER_FALLBACK_MODELS', 'Optional alternate models', 'Comma-separated model IDs.'],
        ['OPENROUTER_TRANSCRIPTION_MODEL', 'Voice-note transcription model', 'Audio usage may require balance.'],
        ['APP_TIME_ZONE', 'Timezone for task context and alerts', 'Asia/Karachi for this deployment.'],
        ['MIN_REPLY_INTERVAL_MS', 'Per-chat flood guard', '15000 ms gives about four replies/minute/chat.'],
        ['IGNORE_GROUPS', 'Ignore group chats', 'true is recommended.'],
    ], [2150, 3500, 3710])
    doc.add_heading('Google Calendar configuration', level=2)
    add_table(doc, ['Variable', 'Purpose'], [
        ['GOOGLE_CALENDAR_CLIENT_ID', 'OAuth client ID from the same Google Cloud project.'],
        ['GOOGLE_CALENDAR_CLIENT_SECRET', 'OAuth client secret. Keep it private.'],
        ['GOOGLE_CALENDAR_REDIRECT_PORT', 'Local callback port; default 3000.'],
        ['GOOGLE_CALENDAR_REMINDER_MINUTES', 'Popup reminder before the created event; default 30.'],
        ['CALENDAR_TOKEN_ENCRYPTION_KEY', 'Optional dedicated secret to encrypt Calendar tokens at rest.'],
    ], [2900, 6460])
    add_callout(doc, 'Google OAuth note:', 'For private testing, configure an External app in Testing status, add the Calendar owner Gmail as a Test user, enable Google Calendar API, and allow the calendar.events scope. The OAuth approval page must be completed in a browser on the same computer that runs the local callback server.', 'F9E6B3')

    chapter(doc, '4. WhatsApp registration and phone pairing')
    add_figure(doc, ASSETS / 'pairing_flow.png', 'Figure 2. A new session can use a pairing code instead of a terminal QR code.')
    doc.add_heading('Existing session', level=2)
    add_paragraph(doc, 'Start all registered sessions:')
    add_callout(doc, 'Command:', 'npm start', LIGHT_GRAY)
    doc.add_heading('Register by QR code', level=2)
    add_callout(doc, 'Command:', 'npm run register -- rehan', LIGHT_GRAY)
    add_paragraph(doc, 'Scan the terminal QR using WhatsApp > Linked devices > Link a device.')
    doc.add_heading('Register by phone-number pairing code', level=2)
    add_callout(doc, 'Command:', 'npm run register:pair -- advisor 923001234567', LIGHT_GRAY)
    add_paragraph(doc, 'Use country code and digits only: no plus sign and no leading zero. On the target phone, open WhatsApp > Linked devices > Link a device > Link with phone number instead, then enter the code printed by the service.')
    add_callout(doc, 'Security:', 'A pairing code gives access to a WhatsApp linked-device session. Show it only to the phone owner and do not store it in logs or screenshots.', 'F9E6B3')

    chapter(doc, '5. Assistant modes and message behavior')
    doc.add_heading('Away Mode', level=2)
    add_paragraph(doc, 'Away Mode is intended for an owner who is unavailable. It sends quick greeting replies, answers general questions, and otherwise uses the configured assistant prompt to acknowledge messages and extract tasks.')
    doc.add_heading('Advisor Mode', level=2)
    add_paragraph(doc, 'Advisor Mode uses a dedicated prompt. It holds natural conversations, matches English or Roman Urdu, and offers supportive practical advice for everyday, relationship, study, and work topics. It must not present itself as an emergency, medical, legal, or financial professional.')
    add_callout(doc, 'Enable Advisor Mode:', 'npm run profile:advisor -- rehan', LIGHT_GRAY)
    doc.add_heading('Owner self-chat commands', level=2)
    add_table(doc, ['Command', 'Effect'], [
        ['!away on', 'Enable automatic replies for this session.'],
        ['!away off', 'Disable automatic replies while retaining logging and owner alerts.'],
        ['!away status', 'Report whether automatic replies are enabled.'],
        ['!tasks', 'List pending saved tasks for the current userId.'],
        ['!calendar connect', 'Start Google OAuth and send the approval URL to owner self-chat.'],
        ['!calendar status', 'Check whether the per-session Calendar token is working.'],
        ['!calendar add 1', 'Manually add an older pending task to Calendar.'],
        ['!calendar disconnect', 'Remove the saved Google Calendar connection for this session.'],
    ], [2700, 6660])
    add_figure(doc, ASSETS / 'message_flow.png', 'Figure 3. The standard incoming-message path, including logging, alerting, and reply selection.')

    chapter(doc, '6. Meeting, tasks, location, and Calendar automation')
    add_figure(doc, ASSETS / 'calendar_flow.png', 'Figure 4. Valid extracted meetings are added to the connected primary Google Calendar.')
    doc.add_heading('Meeting alert contents', level=2)
    add_bullets(doc, ['Place or Google Maps link when the sender explicitly shared a location.', 'Detected time/date reference.', 'Venue when available.', 'Sender display name and WhatsApp chat identifier.', 'Task description, received timestamp, timezone, and original message.'])
    doc.add_heading('Automatic Calendar event behavior', level=2)
    add_numbered(doc, [
        'The model extracts a task marked as a meeting with a valid scheduledFor date and time.',
        'The task is stored under the user session database and receives a unique task ID.',
        'The service checks that this user session has a Google Calendar token.',
        'An event is inserted into that user account primary Calendar with a one-hour duration and configured popup reminder.',
        'The event ID and link are stored to prevent re-adding the same saved task.',
        'The owner self-chat receives a Calendar confirmation and event link.',
    ])
    add_callout(doc, 'Important limitation:', 'If the date/time is unclear, the task is kept locally for review. If the Calendar is disconnected, task saving and WhatsApp replies continue; Calendar creation is skipped and logged.', 'F9E6B3')

    chapter(doc, '7. AI, live information, and voice notes')
    doc.add_heading('General and advice replies', level=2)
    add_paragraph(doc, 'OpenRouter is used for structured JSON replies and task extraction. The reply is sent only after the per-chat flood guard permits it. Common Salaam and "how are you" messages use immediate predefined replies without an AI request.')
    doc.add_heading('Live/current questions', level=2)
    add_paragraph(doc, 'Current questions such as weather, gold/silver price, exchange rates, news, sports scores, availability, and bank timings use the configured live model with web search. The service should not invent a current value; it returns a fallback if no verified answer is available.')
    doc.add_heading('Voice notes', level=2)
    add_paragraph(doc, 'A WhatsApp push-to-talk voice note is downloaded and sent to the transcription endpoint. Its transcript follows the same logging, alert, task, Calendar, and reply workflow as text.')
    add_callout(doc, 'Cost and availability:', 'A 402 audio error means the OpenRouter account does not meet the required audio balance. A 429 error means a provider/model is temporarily rate limited. These are provider-account conditions rather than WhatsApp connection failures.', 'F9E6B3')

    chapter(doc, '8. Operations, monitoring, and troubleshooting')
    doc.add_heading('Routine commands', level=2)
    add_table(doc, ['Command', 'Use'], [
        ['npm start', 'Start every registered session.'],
        ['npm run list-sessions', 'List registered user sessions.'],
        ['npm run calendar:status -- rehan', 'Check Calendar configuration and connection for a user.'],
        ['npm.cmd exec pm2 status', 'Check PM2 service state when deployed with PM2.'],
        ['npm.cmd exec pm2 logs whatsapp-assistant', 'View bot logs.'],
        ['npm.cmd exec pm2 restart whatsapp-assistant', 'Restart the PM2-managed process.'],
        ['npm.cmd exec pm2 stop whatsapp-assistant', 'Stop the PM2-managed process.'],
    ], [3450, 5910])
    doc.add_heading('Common problems', level=2)
    add_table(doc, ['Symptom', 'Likely cause', 'Action'], [
        ['No AI reply', 'Away Mode off or flood guard active', 'Send !away status; wait 15 seconds between messages from the same chat.'],
        ['Delayed / fallback reply', 'Free model 429 rate limit', 'Wait for retry, configure a fallback, or use a reliable paid model.'],
        ['Voice note failed', 'OpenRouter audio 402 or transcription issue', 'Add account credit and check transcription model.'],
        ['Calendar not connected', 'OAuth approval incomplete, test user missing, or callback opened on phone', 'Use browser on the bot computer; check !calendar status.'],
        ['127.0.0.1 callback error on mobile', 'Phone cannot reach computer localhost', 'Open approval link on the same computer as npm start.'],
        ['Session logged out', 'WhatsApp linked device revoked', 'Relink that session with QR or pairing code.'],
    ], [2100, 2750, 4510])

    chapter(doc, '9. Security, privacy, and compliance')
    doc.add_heading('Data stored', level=2)
    add_bullets(doc, [
        'Contact WhatsApp JID, display name (pushName), and message timestamp.',
        'Latitude and longitude only when the contact explicitly sends a WhatsApp location message.',
        'Saved tasks, message deduplication IDs, and per-chat reply timestamps.',
        'Encrypted Google Calendar OAuth tokens and Calendar event mappings when Calendar is connected.',
    ])
    doc.add_heading('Data not collected by this project', level=2)
    add_bullets(doc, ['Device IP addresses.', 'Contact device model details.', 'Location that was not explicitly shared through a WhatsApp location message.'])
    doc.add_heading('Required protections', level=2)
    add_bullets(doc, [
        'Never commit .env, sessions/, WhatsApp credentials, assistant.sqlite, or API keys.',
        'If a key was exposed, revoke it at the provider and create a replacement.',
        'Display a clear contact privacy notice and ensure you have a lawful basis for logging names, JIDs, timestamps, and shared locations.',
        'Use a dedicated CALENDAR_TOKEN_ENCRYPTION_KEY in production rather than relying on an API-key fallback.',
        'For a public dashboard, require user authentication, HTTPS, session ownership checks, rate limits, and audit logs.',
    ])

    chapter(doc, '10. Future web-platform roadmap')
    add_paragraph(doc, 'The current project is a service operated through terminal commands. To let many users link and manage assistants without terminal access, build a secure web platform around the existing session manager.')
    add_numbered(doc, [
        'Add an authenticated web backend with routes for users, sessions, pairing codes, QR display, status, tasks, and Calendar connections.',
        'Build a dashboard where each customer creates an assistant, enters a phone number, receives a pairing code, and sees status.',
        'Enforce ownership so a user can only access their own WhatsApp session, data, and Calendar token.',
        'Move multi-user metadata and logs from local-only storage to a managed database such as PostgreSQL for hosted deployments.',
        'Keep Baileys auth credentials encrypted at rest and use a background worker/process supervisor for every session.',
        'Deploy behind HTTPS with backups, monitoring, abuse controls, and reliable paid AI providers.',
    ])
    add_callout(doc, 'Platform boundary:', 'A browser dashboard removes terminal access for customers, but the server must still run continuously in a protected environment. Pairing codes and WhatsApp credentials remain highly sensitive.', 'F9E6B3')

    chapter(doc, 'Appendix A. File and data layout')
    add_table(doc, ['Path', 'Purpose'], [
        ['index.js', 'Application entry point and multi-session WhatsApp event handler.'],
        ['src/config.js', 'Environment variable parsing and runtime configuration.'],
        ['src/openrouterClient.js', 'AI, web-grounded live answers, and audio transcription.'],
        ['src/dataStore.js', 'Per-user SQLite schema and persistent data access.'],
        ['src/googleCalendar.js', 'Google OAuth token handling and Calendar events.'],
        ['src/sessionStore.js', 'Session paths and per-user metadata.'],
        ['sessions/<userId>/baileys_auth_info/', 'WhatsApp linked-device credentials.'],
        ['sessions/<userId>/assistant.sqlite', 'Per-user contacts, tasks, state, tokens, and event mappings.'],
        ['.env', 'Private runtime configuration and secrets; never commit.'],
        ['README.md', 'Quick-start and operator reference.'],
    ], [3500, 5860])
    chapter(doc, 'Appendix B. Operational checklist')
    add_bullets(doc, [
        'Verify .env includes a valid OpenRouter key and Asia/Karachi timezone if intended.',
        'Run npm.cmd install after pulling dependency changes.',
        'Run npm start and confirm each session logs "WhatsApp user session connected".',
        'For Calendar, complete !calendar connect from a browser on the service computer and confirm !calendar status.',
        'Test greeting, Advisor conversation, valid meeting, live question, and voice note separately.',
        'Review logs for 429 rate limits, 402 audio balance errors, and connection-close events.',
        'Back up session credentials and SQLite databases securely; do not expose them in Git or chat.',
    ])

    chapter(doc, 'Appendix C. New amendment page template')
    add_paragraph(doc, 'Copy or print this page whenever a new feature, command, workflow, or configuration item is added. Insert the completed page immediately after the related numbered module.')
    add_table(doc, ['Update field', 'Complete this when amending the manual'], [
        ['Amendment date', ''],
        ['Related module number', ''],
        ['Feature or configuration changed', ''],
        ['Code files changed', ''],
        ['Environment variables changed', ''],
        ['New command(s)', ''],
        ['Workflow impact', ''],
        ['Testing completed', ''],
        ['Printed pages to replace / insert', ''],
        ['Approved by', ''],
    ], [2900, 6460])
    doc.add_heading('Change description', level=2)
    for _ in range(8):
        p = doc.add_paragraph('________________________________________________________________________________')
        p.paragraph_format.space_after = Pt(8)
        for run in p.runs:
            set_run(run, 10, MUTED)

    doc.save(OUT)


if __name__ == '__main__':
    build_document()
    print(OUT)
